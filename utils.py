import pandas as pd
from docx import Document

# 1. Validate Excel File
import pandas as pd

def validate_excel(file_path):
    expected_sheets = {'Transactions', 'Customers', 'Products'}
    expected_columns = {
        'Transactions': {'transaction_id', 'customer_id', 'transaction_date', 'product_code', 'amount', 'payment_type'},
        'Customers': {0}, 
        'Products': {'product_code', 'product_name', 'category', 'unit_price'}
    }

    xls = pd.ExcelFile(file_path)
    found_sheets = set(xls.sheet_names)

    # 1. Check sheets exist
    missing_sheets = expected_sheets - found_sheets
    if missing_sheets:
        raise ValueError(f"Missing sheets: {missing_sheets}")

    # 2. Check Transactions and Products columns
    for sheet in ['Transactions', 'Products']:
        df = pd.read_excel(xls, sheet_name=sheet)
        missing_cols = expected_columns[sheet] - set(df.columns)
        if missing_cols:
            raise ValueError(f"Sheet '{sheet}' is missing columns: {missing_cols}")


    df_customers_raw = pd.read_excel(xls, sheet_name='Customers', header=None)
    if df_customers_raw.empty:
        raise ValueError("Customers sheet is empty or incorrect format.")

    # 4. Check minimum data presence for Transactions
    df_trans = pd.read_excel(xls, sheet_name='Transactions')
    if df_trans['customer_id'].isnull().any():
        raise ValueError("Transactions sheet has missing customer_id values.")

    return True


# 2. Load Data
def parse_customers(raw_df):
    customer_rows = raw_df.iloc[:, 0]
    parsed = []

    for row in customer_rows:
        clean = row.strip('{}')
        parts = clean.split('_')
        if len(parts) == 6:
            parsed.append(parts)

    df = pd.DataFrame(parsed, columns=[
        'customer_id', 'name', 'email', 'dob', 'address', 'created_date'
    ])
    df['created_date'] = pd.to_numeric(df['created_date'], errors='coerce')
    df['created_date'] = pd.to_datetime(df['created_date'], unit='D', origin='1899-12-30')
    return df

def load_data(file_path):
    xls = pd.ExcelFile(file_path)
    trans = pd.read_excel(xls, 'Transactions')
    products = pd.read_excel(xls, 'Products')
    raw_customers = pd.read_excel(xls, 'Customers', header=None)
    customers = parse_customers(raw_customers)
    return trans, customers, products

# 3a. Address Changes
def detect_address_changes(customers_df):
    customers_df = customers_df.sort_values(['customer_id', 'created_date'])
    customers_df['prev_address'] = customers_df.groupby('customer_id')['address'].shift()
    print(customers_df)
    customers_df['address_changed'] = customers_df['address'] != customers_df['prev_address'] 
    print(customers_df)
    return customers_df[customers_df['address_changed']]

# 3b. Total Spend per Category
def transaction_summary(trans_df, products_df):
    merged = trans_df.merge(products_df[['product_code', 'category']], on='product_code')
    return merged.groupby(['customer_id', 'category'])['amount'].sum().reset_index()

# 3c. Top Spenders
def top_spenders(summary_df):
    return summary_df.loc[summary_df.groupby('category')['amount'].idxmax()]

# 3d. Customer Ranking
def rank_customers(trans_df):
    total = trans_df.groupby('customer_id')['amount'].sum().reset_index()
    total['rank'] = total['amount'].rank(ascending=False, method='min').astype(int)
    return total.sort_values('rank')

# 6. Save Processed Output
def save_processed_output(**dfs):
    path = 'output/processed_output.xlsx'
    with pd.ExcelWriter(path) as writer:
        for name, df in dfs.items():
            df.to_excel(writer, sheet_name=name, index=False)
    return path

# 7. Summary Report
def make_summary_report(summary, top, ranks):
    doc = Document()
    doc.add_heading('Data Summary Report', 0)

    doc.add_paragraph("✅ Total transaction summary per customer per category:")
    for i, row in summary.iterrows():
        doc.add_paragraph(f"{row['customer_id']} - {row['category']}: ${row['amount']:.2f}")

    doc.add_paragraph("\n🏆 Top Spenders:")
    for i, row in top.iterrows():
        doc.add_paragraph(f"{row['category']}: {row['customer_id']} (${row['amount']:.2f})")

    doc.add_paragraph("\n📊 Top Customer Rankings:")
    for i, row in ranks.iterrows():
        doc.add_paragraph(f"Rank {row['rank']}: {row['customer_id']} (${row['amount']:.2f})")

    summary_path = 'output/summary_report.docx'
    doc.save(summary_path)
    return summary_path
